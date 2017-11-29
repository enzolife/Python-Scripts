import urllib.request
from bs4 import BeautifulSoup
from docx import Document
from xhtml2pdf import pisa


html = urllib.request.urlopen('http://www.woshipm.com/data-analysis/735340.html')
bsObj = BeautifulSoup(html.read())
article = bsObj.find('div', {'class': 'entry-content'})


# Define your data
sourceHtml = "<html><body><p>To PDF or not to PDF</p></body></html>"
outputFilename = "test.pdf"


# Utility function
def convertHtmlToPdf(sourceHtml, outputFilename):
    # open output file for writing (truncated binary)
    resultFile = open(outputFilename, "w+b")

    # convert HTML to PDF
    pisaStatus = pisa.CreatePDF(
            sourceHtml,                # the HTML to convert
            dest=resultFile)           # file handle to recieve result

    # close output file
    resultFile.close()                 # close output file

    # return True on success and False on errors
    return pisaStatus.err

# Main program
if __name__ == "__main__":
    pisa.showLogging()
    convertHtmlToPdf(sourceHtml, outputFilename)




'''
document = Document()

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    article.get_text(), style='ListBullet'
)

document.save('demo.docx')
'''